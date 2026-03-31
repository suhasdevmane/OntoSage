"""
Document Builder — PDF/DOCX/HTML report generation via Jinja2 templates
=========================================================================
Renders structured report data into professional HTML documents, then
converts to PDF (via pdfkit/wkhtmltopdf) or DOCX (via python-docx).

Template selection is driven by (report_type, persona) mapping.

Usage:
    from orchestrator.services.document_builder import DocumentBuilder

    builder = DocumentBuilder()
    result = builder.render(
        report_data={"narrative": "...", "readings": {...}, ...},
        report_type="summary",
        persona="executive",
        output_format="pdf",
        title="Weekly Building KPI Report",
    )
    # result = {"success": True, "content": <bytes/str>, "filename": "...", ...}
"""
from __future__ import annotations

import io
import logging
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
from zoneinfo import ZoneInfo

from shared.config import settings

logger = logging.getLogger(__name__)

_TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "document_templates"

# Map (report_type, persona) -> template filename
# Falls back to report_type-only match, then to weekly_summary.html
TEMPLATE_MAP = {
    ("summary", "facility_manager"): "weekly_summary.html",
    ("summary", "executive"): "executive_kpi.html",
    ("summary", "general"): "weekly_summary.html",
    ("anomaly", "safety_officer"): "anomaly_digest.html",
    ("anomaly", "facility_manager"): "anomaly_digest.html",
    ("anomaly", "general"): "anomaly_digest.html",
    ("compliance", "sustainability_officer"): "compliance_report.html",
    ("compliance", "general"): "compliance_report.html",
    ("full", "executive"): "executive_kpi.html",
    ("full", "general"): "weekly_summary.html",
    ("trend", "general"): "weekly_summary.html",
    ("comparison", "general"): "weekly_summary.html",
}

# Fallback per report_type (persona-independent)
_TYPE_FALLBACK = {
    "summary": "weekly_summary.html",
    "anomaly": "anomaly_digest.html",
    "compliance": "compliance_report.html",
    "trend": "weekly_summary.html",
    "comparison": "weekly_summary.html",
    "full": "weekly_summary.html",
}


def _select_template(report_type: str, persona: str) -> str:
    return (
        TEMPLATE_MAP.get((report_type, persona))
        or _TYPE_FALLBACK.get(report_type)
        or "weekly_summary.html"
    )


class DocumentBuilder:
    """Renders report data into HTML/PDF/DOCX documents."""

    def __init__(self, template_dir: Optional[Path] = None):
        self._template_dir = template_dir or _TEMPLATE_DIR
        self._jinja_env = None

    def _get_jinja_env(self):
        if self._jinja_env is None:
            from jinja2 import Environment, FileSystemLoader, select_autoescape
            self._jinja_env = Environment(
                loader=FileSystemLoader(str(self._template_dir)),
                autoescape=select_autoescape(["html"]),
            )
        return self._jinja_env

    def render(
        self,
        report_data: Dict[str, Any],
        report_type: str = "summary",
        persona: str = "general",
        output_format: str = "html",
        title: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Render a report to the specified format.

        Args:
            report_data: Dict with keys like narrative, readings, anomalies, highlights, kpis, etc.
            report_type: summary / anomaly / compliance / trend / comparison / full
            persona: User persona for template selection
            output_format: html / pdf / docx
            title: Report title

        Returns:
            Dict with success, content (str or bytes), filename, format, size_bytes
        """
        output_format = output_format.lower().strip()
        if output_format not in ("html", "pdf", "docx"):
            return {"success": False, "error": f"Unsupported document format: {output_format}"}

        try:
            template_name = _select_template(report_type, persona)
            logger.info(f"DocumentBuilder: template={template_name} format={output_format}")

            # Build template context
            tz = ZoneInfo(settings.BUILDING_TIMEZONE) if hasattr(settings, "BUILDING_TIMEZONE") else None
            generated_at = datetime.now(tz).strftime("%Y-%m-%d %H:%M") if tz else datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")
            building_name = getattr(settings, "BUILDING_NAME", "Smart Building")

            ctx = {
                "title": title or f"{report_type.title()} Report",
                "building_name": building_name,
                "generated_at": generated_at,
                "report_type": report_type,
                "narrative": report_data.get("narrative") or report_data.get("formatted_text") or "",
                "readings": report_data.get("readings_summary") or report_data.get("readings") or {},
                "anomalies": report_data.get("anomalies") or [],
                "highlights": report_data.get("highlights") or [],
                "kpis": report_data.get("kpis") or self._auto_kpis(report_data),
                "standards_results": report_data.get("standards_results") or [],
            }

            # Render HTML
            env = self._get_jinja_env()
            template = env.get_template(template_name)
            html_content = template.render(**ctx)

            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_filename = f"report_{report_type}_{ts}"

            if output_format == "html":
                return self._result(html_content, f"{base_filename}.html", "html")

            if output_format == "pdf":
                return self._render_pdf(html_content, f"{base_filename}.pdf")

            if output_format == "docx":
                return self._render_docx(report_data, ctx, f"{base_filename}.docx")

        except Exception as e:
            logger.error(f"DocumentBuilder render failed: {e}", exc_info=True)
            return {"success": False, "error": str(e)}

    def _auto_kpis(self, report_data: Dict) -> List[Dict]:
        """Auto-generate KPI cards from readings summary."""
        kpis = []
        readings = report_data.get("readings_summary") or report_data.get("readings") or {}
        for name, stats in list(readings.items())[:6]:
            if isinstance(stats, dict) and "avg" in stats:
                kpis.append({"label": name, "value": stats["avg"], "unit": ""})
        anomalies = report_data.get("anomalies") or []
        if anomalies:
            kpis.append({"label": "Anomalies", "value": len(anomalies), "unit": "events"})
        return kpis

    def _render_pdf(self, html_content: str, filename: str) -> Dict[str, Any]:
        """Convert HTML to PDF using pdfkit (wkhtmltopdf)."""
        try:
            import pdfkit
            pdf_bytes = pdfkit.from_string(html_content, False, options={
                "page-size": "A4",
                "margin-top": "15mm",
                "margin-bottom": "15mm",
                "margin-left": "15mm",
                "margin-right": "15mm",
                "encoding": "UTF-8",
                "quiet": "",
            })
            return self._result(pdf_bytes, filename, "pdf")
        except ImportError:
            logger.warning("pdfkit not installed — falling back to HTML output")
            html_filename = filename.replace(".pdf", ".html")
            result = self._result(html_content, html_filename, "html")
            result["fallback"] = True
            result["fallback_reason"] = "pdfkit not installed; install pdfkit + wkhtmltopdf for PDF"
            return result
        except Exception as e:
            logger.warning(f"PDF generation failed: {e} — falling back to HTML")
            html_filename = filename.replace(".pdf", ".html")
            result = self._result(html_content, html_filename, "html")
            result["fallback"] = True
            result["fallback_reason"] = str(e)
            return result

    def _render_docx(self, report_data: Dict, ctx: Dict, filename: str) -> Dict[str, Any]:
        """Generate a Word document using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title_para = doc.add_heading(ctx["title"], level=0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            meta = doc.add_paragraph()
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = meta.add_run(f"{ctx['building_name']} | {ctx['generated_at']} | {ctx['report_type'].title()} Report")
            run.font.size = Pt(10)

            # Narrative
            if ctx.get("narrative"):
                doc.add_heading("Executive Summary", level=1)
                doc.add_paragraph(ctx["narrative"])

            # Readings table
            readings = ctx.get("readings") or {}
            if readings:
                doc.add_heading("Sensor Readings", level=1)
                table = doc.add_table(rows=1, cols=5)
                table.style = "Light Grid Accent 1"
                hdr = table.rows[0].cells
                for i, h in enumerate(["Metric", "Count", "Min", "Avg", "Max"]):
                    hdr[i].text = h
                for name, stats in readings.items():
                    if isinstance(stats, dict):
                        row = table.add_row().cells
                        row[0].text = str(name)
                        row[1].text = str(stats.get("count", ""))
                        row[2].text = str(stats.get("min", ""))
                        row[3].text = str(stats.get("avg", ""))
                        row[4].text = str(stats.get("max", ""))

            # Anomalies
            anomalies = ctx.get("anomalies") or []
            if anomalies:
                doc.add_heading(f"Anomalies ({len(anomalies)})", level=1)
                table = doc.add_table(rows=1, cols=4)
                table.style = "Light Grid Accent 1"
                hdr = table.rows[0].cells
                for i, h in enumerate(["Sensor", "Value", "Severity", "Timestamp"]):
                    hdr[i].text = h
                for a in anomalies[:30]:
                    row = table.add_row().cells
                    row[0].text = str(a.get("column", ""))
                    row[1].text = f"{a.get('value', '')} {a.get('unit', '')}"
                    row[2].text = str(a.get("severity", "")).upper()
                    row[3].text = str(a.get("timestamp", ""))

            # Highlights
            highlights = ctx.get("highlights") or []
            if highlights:
                doc.add_heading("Key Findings", level=1)
                for h in highlights:
                    doc.add_paragraph(h, style="List Bullet")

            # Footer
            doc.add_paragraph("")
            footer = doc.add_paragraph()
            run = footer.add_run(f"Generated by OntoSage | {ctx['building_name']} | {ctx['generated_at']}")
            run.font.size = Pt(8)

            buf = io.BytesIO()
            doc.save(buf)
            docx_bytes = buf.getvalue()
            return self._result(docx_bytes, filename, "docx")

        except ImportError:
            logger.warning("python-docx not installed — falling back to HTML")
            return {"success": False, "error": "python-docx not installed; install it for DOCX output", "fallback": True}

    def _result(self, content, filename: str, fmt: str) -> Dict[str, Any]:
        size = len(content) if isinstance(content, (bytes, str)) else 0
        return {
            "success": True,
            "format": fmt,
            "filename": filename,
            "content": content,
            "size_bytes": size,
        }

    def save_to_exports(self, result: Dict[str, Any]) -> Optional[str]:
        """Save rendered document to EXPORTS_DIR and return download URL."""
        if not result.get("success") or not result.get("content"):
            return None
        try:
            exports_dir = Path(settings.EXPORTS_DIR)
            exports_dir.mkdir(parents=True, exist_ok=True)
            file_path = exports_dir / result["filename"]
            content = result["content"]
            if isinstance(content, bytes):
                file_path.write_bytes(content)
            else:
                file_path.write_text(content, encoding="utf-8")
            url = f"{settings.STATIC_BASE_URL}/static/exports/{result['filename']}"
            logger.info(f"DocumentBuilder: saved to {file_path} -> {url}")
            return url
        except Exception as e:
            logger.warning(f"DocumentBuilder save failed: {e}")
            return None
